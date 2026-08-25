"""Document ingest: extract -> (OCR when scanned) -> clean -> chunk."""
from __future__ import annotations

import base64
import io
import logging
import re
from typing import Any

import httpx

from .. import config
from ..models.engine import embedder

log = logging.getLogger("rag.ingest")

TEXT_PER_PAGE_MIN = 60  # fewer chars on a PDF page => treat as scanned


class IngestError(Exception):
    pass


def sniff_type(filename: str, head: bytes) -> str:
    if head.startswith(b"%PDF"):
        return "pdf"
    if head.startswith(b"PK\x03\x04") and filename.lower().endswith(".docx"):
        return "docx"
    if filename.lower().endswith(".txt"):
        return "txt"
    raise IngestError("Unsupported file. Use PDF, DOCX or TXT.")


def extract_pages(path: str, kind: str) -> list[tuple[int, str]]:
    """Returns [(page_number, text)]. Pages with no text layer come back empty (OCR candidates)."""
    if kind == "pdf":
        import pdfplumber

        pages: list[tuple[int, str]] = []
        with pdfplumber.open(path) as pdf:
            for i, page in enumerate(pdf.pages, start=1):
                text = page.extract_text() or ""
                pages.append((i, clean(text)))
        return pages
    if kind == "docx":
        import docx

        d = docx.Document(path)
        text = clean("\n".join(p.text for p in d.paragraphs))
        return [(1, text)]  # DOCX has no page model: single logical unit
    with open(path, encoding="utf-8", errors="replace") as f:
        return [(1, clean(f.read()))]


def clean(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\xa0", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def needs_ocr(pages: list[tuple[int, str]], kind: str) -> list[int]:
    if kind != "pdf":
        return []
    return [n for n, t in pages if len(t) < TEXT_PER_PAGE_MIN]


async def ocr_pages(path: str, page_numbers: list[int]) -> dict[int, str]:
    """zerox approach: render page -> PNG -> Gemini vision -> markdown text.
    Gemini only: opencode zen (MiMo) has no vision, so OCR is unavailable under LLM_PROVIDER=zen."""
    if not page_numbers:
        return {}
    if config.LLM_PROVIDER != "gemini" or not config.GEMINI_API_KEY:
        raise IngestError("This PDF looks scanned and OCR is not configured (Gemini vision required).")

    import pypdfium2 as pdfium

    out: dict[int, str] = {}
    async with httpx.AsyncClient(timeout=120) as client:
        pdf = pdfium.PdfDocument(path)
        try:
            for n in page_numbers:
                page = pdf[n - 1]
                bitmap = page.render(scale=150 / 72)
                pil_image = bitmap.to_pil()
                buf = io.BytesIO()
                pil_image.save(buf, format="PNG")
                png_b64 = base64.b64encode(buf.getvalue()).decode()
                body = {
                    "contents": [{"parts": [
                        {"text": "Transcribe all text from this page as plain markdown. Output text only, no commentary."},
                        {"inline_data": {"mime_type": "image/png", "data": png_b64}},
                    ]}]
                }
                url = (f"https://generativelanguage.googleapis.com/v1beta/models/"
                       f"{config.GEMINI_OCR_MODEL}:generateContent?key={config.GEMINI_API_KEY}")
                r = await client.post(url, json=body)
                if r.status_code == 429:
                    raise IngestError("Daily OCR quota reached. Try again tomorrow or upload a text-based PDF.")
                r.raise_for_status()
                data = r.json()
                text = "".join(
                    p.get("text", "")
                    for c in data.get("candidates", [])
                    for p in c.get("content", {}).get("parts", [])
                )
                out[n] = clean(text)
        finally:
            pdf.close()
    return out


def chunk_pages(pages: list[tuple[int, str]]) -> list[tuple[int, int, str]]:
    """Chunks bounded by the e5-small 512-token limit (incl. the 'passage: ' prefix).
    Never truncate silently: the chunker splits before the limit is reached."""
    chunks: list[tuple[int, int, str]] = []
    idx = 0
    for page, text in pages:
        if not text.strip():
            continue
        for body in _split_tokens(text):
            chunks.append((page, idx, body))
            idx += 1
    return chunks


def _split_tokens(text: str) -> list[str]:
    paragraphs = [p.strip() for p in re.split(r"\n{2,}|(?<=[.!?])\s+(?=[A-ZА-Я])", text) if p.strip()]
    out: list[str] = []
    cur: list[str] = []
    cur_tokens = 0
    for p in paragraphs:
        n = embedder.count_tokens(p)
        if n > config.CHUNK_TOKENS:  # oversized paragraph -> hard split by words
            for piece in _split_words(p):
                out.append(piece)
            continue
        if cur_tokens + n > config.CHUNK_TOKENS and cur:
            out.append(" ".join(cur))
            overlap: list[str] = []
            otok = 0
            for part in reversed(cur):
                t = embedder.count_tokens(part)
                if otok + t > config.CHUNK_OVERLAP:
                    break
                overlap.insert(0, part)
                otok += t
            cur, cur_tokens = overlap, otok
        cur.append(p)
        cur_tokens += n
    if cur:
        out.append(" ".join(cur))
    return [c for c in out if c.strip()]


def _split_words(paragraph: str) -> list[str]:
    words = paragraph.split()
    out: list[str] = []
    cur: list[str] = []
    for w in words:
        cur.append(w)
        if len(cur) % 40 == 0 and embedder.count_tokens(" ".join(cur)) > config.CHUNK_TOKENS - 20:
            out.append(" ".join(cur))
            cur = cur[-10:]  # small word overlap
    if cur:
        out.append(" ".join(cur))
    return out
